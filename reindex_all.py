#!/usr/bin/env python3
"""
NURU — Réindexation complète des documents utilisateur.

Scanne ~/Documents, ~/Desktop, ~/Downloads pour les fichiers
.pdf, .docx, .txt, .md, .csv, .json.

Étapes pour chaque fichier nouveau ou modifié :
  1. IngestionEngine.index_file()  → met à jour l'index vectoriel
  2. extract_document()             → produit des métadonnées structurées
  3. Sauvegarde dans doc_structured / cv_structured

Respecte la RAM : si moins de 500 Mo libre, pause de 30 s entre les fichiers.

Usage :
   cd /Users/leblancbahiga/Downloads/Assistant_IA    # (ou le chemin réel)
   .venv/bin/python3 reindex_all.py
"""

import asyncio
import hashlib
import logging
import os
import sys
import time
from pathlib import Path

# ── Projet ──
# Ajouter src/ au path pour importer les modules du projet
PROJECT_ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(PROJECT_ROOT))

try:
    import psutil
except ImportError:
    psutil = None  # RAM check désactivé si psutil n'est pas installé

from src.ingestion import IngestionEngine
from src.rag_engine import RAGEngine
from src.document_extractor import extract_document

# ── Configuration ──
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json"}
SCAN_DIRS = [
    Path.home() / "Documents",
    Path.home() / "Desktop",
    Path.home() / "Downloads",
]
RAM_PAUSE_THRESHOLD_MB = 500
RAM_PAUSE_SECONDS = 30
LOG_FORMAT = "%(asctime)s [%(levelname)s] %(message)s"

logging.basicConfig(
    level=logging.INFO,
    format=LOG_FORMAT,
    datefmt="%H:%M:%S",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger("reindex")


# ── Utilitaires ──

def compute_sha256(filepath: str) -> str:
    """Calcule le hash SHA256 d'un fichier."""
    hasher = hashlib.sha256()
    try:
        with open(filepath, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()
    except Exception as e:
        logger.warning("⚠️  Impossible de hacher %s : %s", filepath, e)
        return ""


def check_ram_mb() -> float:
    """Retourne la RAM libre disponible en Mo (ou un grand nombre si psutil absent)."""
    if psutil is None:
        return 9999.0  # pas de limitation
    return psutil.virtual_memory().available / (1024 * 1024)


async def pause_if_low_ram():
    """Pause 30 s si la RAM libre passe sous le seuil critique."""
    free_mb = check_ram_mb()
    if free_mb < RAM_PAUSE_THRESHOLD_MB:
        logger.warning(
            "🧠  RAM faible : %.0f Mo libre — pause %d s",
            free_mb,
            RAM_PAUSE_SECONDS,
        )
        await asyncio.sleep(RAM_PAUSE_SECONDS)
        # Re-vérifier après la pause
        free_mb = check_ram_mb()
        logger.info("🧠  RAM après pause : %.0f Mo libre", free_mb)


def extract_text_from_file(filepath: str) -> str:
    """
    Extrait le texte brut d'un fichier selon son extension.
    (Copie légère de IngestionEngine._parse_file pour pouvoir passer
    le texte à extract_document() sans ré-indexer.)
    """
    path = Path(filepath)
    ext = path.suffix.lower()
    text = ""

    try:
        if ext == ".pdf":
            import fitz

            with fitz.open(str(path)) as doc:
                for page in doc:
                    text += page.get_text()
        elif ext == ".docx":
            from docx import Document

            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(
                        cell.text for cell in row.cells if cell.text.strip()
                    )
                    if row_text.strip():
                        parts.append(row_text)
            text = "\n".join(parts)
        elif ext in (".txt", ".md"):
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
        elif ext == ".csv":
            import csv

            with open(path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                text = "\n".join([",".join(row) for row in reader])
        elif ext == ".json":
            import json as j

            with open(path, "r", encoding="utf-8") as f:
                data = j.load(f)
                text = j.dumps(data, indent=2, ensure_ascii=False)
    except Exception as e:
        logger.error("❌ Erreur d'extraction de texte pour %s : %s", filepath, e)

    return text


# ── Logique principale ──

async def process_file(
    filepath: str,
    ingestion: IngestionEngine,
    rag: RAGEngine,
) -> bool:
    """
    Traite un fichier unique :
      1. Vérifie si déjà indexé (hash)
      2. Lance index_file()
      3. Extrait le texte et appelle extract_document()
      4. Sauvegarde les métadonnées structurées
    Retourne True si le fichier a été (ré)indexé.
    """
    # 1. Vérification par hash
    file_hash = compute_sha256(filepath)
    if not file_hash:
        return False

    basename = os.path.basename(filepath)

    if rag.is_file_up_to_date(filepath, 0, file_hash):
        logger.info("⏭️  Déjà indexé : %s", basename)
        return False

    # 2. Indexation vectorielle
    logger.info("📄 Indexation : %s", basename)
    await ingestion.index_file(filepath)

    # Vérifier que l'indexation a bien marqué le fichier
    if rag.is_file_up_to_date(filepath, 0, file_hash):
        logger.info("✅  Indexation OK : %s", basename)
    else:
        logger.warning("⚠️  index_file() n'a pas marqué %s comme indexé", basename)
        # On continue quand même avec l'extraction structurée

    # 3. Extraction structurée (extract_document)
    await pause_if_low_ram()

    logger.info("🔍  Extraction structurée : %s", basename)
    text = extract_text_from_file(filepath)
    if not text:
        logger.warning("⚠️  Aucun texte extrait pour %s — skip extraction", basename)
        return True  # indexé mais pas de métadonnées

    meta = await extract_document(text, basename)
    if meta is None:
        logger.warning("⚠️  extract_document() a retourné None pour %s", basename)
        return True

    # 4. Sauvegarde dans la table appropriée
    source_name = basename
    if meta.doc_type == "CV":
        rag.save_cv(source_name, file_hash, meta.structured_json)
        logger.info("📋  CV structuré sauvegardé : %s", basename)
    else:
        rag.save_doc_meta(
            source_name,
            file_hash,
            meta.doc_type,
            meta.structured_json,
        )
        logger.info(
            "📋  Métadonnées sauvegardées : %s (%s)",
            basename,
            meta.doc_type,
        )

    return True


async def main():
    """Point d'entrée : scanne, indexe, extrait."""
    start_time = time.time()
    logger.info("=" * 60)
    logger.info("🚀  RÉINDEXATION COMPLÈTE — DÉMARRAGE")
    logger.info("=" * 60)

    # Chemins
    for d in SCAN_DIRS:
        if d.exists():
            logger.info("📁  Scan : %s", d)
        else:
            logger.warning("📁  Introuvable : %s — ignoré", d)

    # Instances
    ingestion = IngestionEngine()
    rag = RAGEngine()

    total_indexed = 0
    total_skipped = 0
    total_errors = 0
    files_scanned = 0

    for base_dir in SCAN_DIRS:
        if not base_dir.exists():
            continue

        for root, _, files in os.walk(base_dir):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext not in SUPPORTED_EXTENSIONS:
                    continue

                filepath = os.path.join(root, file)
                files_scanned += 1

                try:
                    ok = await process_file(filepath, ingestion, rag)
                    if ok:
                        total_indexed += 1
                    else:
                        total_skipped += 1
                except Exception as e:
                    logger.error(
                        "❌  Erreur sur %s : %s",
                        os.path.basename(filepath),
                        e,
                    )
                    total_errors += 1

                # Pause RAM entre chaque fichier
                await pause_if_low_ram()

                # Petite pause CPU (identique à auto_index_loop)
                await asyncio.sleep(0.1)

    elapsed = time.time() - start_time
    minutes, secs = divmod(int(elapsed), 60)

    logger.info("=" * 60)
    logger.info("🏁  RÉINDEXATION TERMINÉE")
    logger.info("📊  Bilan :")
    logger.info("      Fichiers scannés     : %d", files_scanned)
    logger.info("      Indexés              : %d", total_indexed)
    logger.info("      Déjà indexés (skipped) : %d", total_skipped)
    logger.info("      Erreurs              : %d", total_errors)
    logger.info("⏱️   Temps total           : %d min %d s", minutes, secs)
    logger.info("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())
