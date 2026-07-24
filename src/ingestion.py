import os
import asyncio
import logging
import hashlib
from pathlib import Path
from typing import List, Optional
import numpy as np
from src.config import config
from src.rag_engine import RAGEngine
from src.embedder import Embedder
from src.document_extractor import extract_document, detect_doc_type

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json"}


class IngestionEngine:
    """Moteur d'ingestion de documents pour le RAG V8+."""
    
    def __init__(self, rag_engine=None):
        self.rag = rag_engine or RAGEngine()  # V16 FIX : accepter RAGEngine existant
        self.embedder = Embedder()

    def _parse_file(self, file_path: str) -> str:
        """Extrait le texte brut d'un fichier selon son extension.
        V6.1 : Fallback OCR pour les PDF scannés via Tesseract.
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        text = ""
        
        try:
            if ext == ".pdf":
                import fitz  # lazy — évite crash si PyMuPDF manquant
                doc = fitz.open(str(path))
                try:
                    page_count = len(doc)
                    for page in doc:
                        text += page.get_text()
                finally:
                    doc.close()
                # V6.1 : Si PyMuPDF n'a rien extrait → PDF scanné → OCR
                if not text.strip() and page_count > 0:
                    from src.ocr import ocr_fallback
                    ocr_text = ocr_fallback(str(path), "")
                    if ocr_text:
                        text = ocr_text
            elif ext == ".docx":
                from docx import Document  # lazy — évite crash si python-docx manquant
                doc = Document(str(path))
                # V12 : extraire aussi le contenu des tables AVEC en-têtes de colonnes
                parts = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    if not table.rows:
                        continue
                    # Extraire les en-têtes de colonnes depuis la première ligne
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    for i, row in enumerate(table.rows):
                        cell_texts = []
                        for j, cell in enumerate(row.cells):
                            val = cell.text.strip()
                            if not val:
                                continue
                            # Préfixer les cellules de données avec leur en-tête de colonne
                            if i > 0 and j < len(headers) and headers[j]:
                                cell_texts.append(f"{headers[j]}: {val}")
                            else:
                                cell_texts.append(val)
                        row_text = "\t".join(cell_texts)
                        if row_text.strip():
                            parts.append(row_text)
                text = "\n".join(parts)
            elif ext in [".txt", ".md"]:
                # V16 FIX : fallback latin-1 si UTF-8 échoue (fichiers Windows)
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        text = f.read()
                except UnicodeDecodeError:
                    with open(path, "r", encoding="latin-1") as f:
                        text = f.read()
                        logger.info(f"  ↪ latin-1 fallback: {path.name}")
            elif ext == ".csv":
                import csv
                import io
                decoded = ""
                # V16 FIX : fallback latin-1 + ignore les lignes malformées
                try:
                    raw = path.read_bytes()
                    for enc in ("utf-8", "latin-1", "cp1252"):
                        try:
                            decoded = raw.decode(enc)
                            break
                        except UnicodeDecodeError:
                            continue
                    if not decoded:
                        decoded = raw.decode("utf-8", errors="replace")
                    reader = csv.reader(io.StringIO(decoded))
                    text = "\n".join([",".join(row) for row in reader if any(cell.strip() for cell in row)])
                except Exception as csv_err:
                    logger.error(f"⚠️ CSV {path.name}: {csv_err}")
                    text = decoded or ""
            elif ext == ".json":
                with open(path, "r", encoding="utf-8") as f:
                    import json as j
                    data = j.load(f)
                    text = j.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Erreur lors du parsing de {file_path} : {e}")
            
        return text

    def compute_sha256(self, filepath: str) -> str:
        """Calcule le hash SHA256 d'un fichier pour détecter les vrais changements."""
        hasher = hashlib.sha256()
        try:
            with open(filepath, 'rb') as f:
                # Lecture par morceaux pour les gros fichiers
                for chunk in iter(lambda: f.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"Erreur hashage {filepath}: {e}")
            return ""

    async def index_file(self, filepath: str):
        """Parse et indexe un fichier unique avec timeout de sécurité."""
        try:
            # V15 Phase 0B : validation path contre Path Traversal
            from src.core.prompt_guard import sanitize_path
            sanitize_path(filepath)

            # Timeout critique de 120s par document (V10.1: 30s → 120s pour les gros fichiers comme NURU_V9.md)
            async with asyncio.timeout(120):
                file_hash = self.compute_sha256(filepath)
                if not file_hash:
                    return

                # On vérifie le hash au lieu de mtime
                if self.rag.is_file_up_to_date(filepath, 0, file_hash):
                    return

                logger.info(f"Indexing: {filepath}")
                text = self._parse_file(filepath)
                if not text:
                    return

                # Phase 2 : Chunking sémantique contextuel
                from src.rag.chunking import SemanticChunker
                chunker = SemanticChunker()
                metadata = {"title": os.path.basename(filepath), "source": filepath}
                semantic_chunks = chunker.chunk(text, source=os.path.basename(filepath),
                                                metadata=metadata)

                # NURU V6 : Chunking hiérarchique V2 (plus robuste, chunks plus gros)
                from src.rag.v2_chunking import HierarchicalChunkerV2
                # V10.1 : detecter le profil adapte au fichier
                profile = HierarchicalChunkerV2.detect_profile(
                    os.path.basename(filepath)
                )
                v2_chunker = HierarchicalChunkerV2(profile=profile)
                v2_chunks = v2_chunker.chunk(
                    text,
                    source=os.path.basename(filepath),
                    doc_title=os.path.basename(filepath),
                )
                
                # Utiliser V2 comme source principale, V1 comme fallback
                primary_chunks = v2_chunks if v2_chunks else semantic_chunks

                chunks = []
                # V15 #72 : Batch Processing — lots de 32 chunks max
                # pour éviter OOM sur M1 8 Go et +2-3x vitesse d'indexation
                chunk_dicts = [c.to_dict() for c in primary_chunks]
                contents = [d["content"] for d in chunk_dicts]
                if contents:
                    batch_size = 32
                    all_embeddings = []
                    for i in range(0, len(contents), batch_size):
                        batch = contents[i:i + batch_size]
                        batch_emb = await self.embedder.embed(batch, is_query=False)
                        if batch_emb is not None and len(batch_emb) > 0:
                            all_embeddings.append(batch_emb)
                        await asyncio.sleep(0.01)  # laisse respirer le CPU
                    all_embeddings = np.concatenate(all_embeddings) if all_embeddings else None
                    for i, c in enumerate(primary_chunks):
                        chunks.append({
                            "content": contents[i],
                            "source": os.path.basename(filepath),
                            "embedding": all_embeddings[i] if all_embeddings is not None and i < len(all_embeddings) else None,
                            "date": "",
                            "title": c.section_title or c.doc_title,
                            "level": c.level,
                        })

                self.rag.add_chunks(chunks)
                self.rag.mark_file_indexed(filepath, 0, file_hash)
                # NURU V6 : Dual-Write — chaque chunk aussi dans ~/Nuru_Brain/
                try:
                    from src.nuru_brain import WikiWriter
                    wiki = WikiWriter()
                    source_name = os.path.basename(filepath)
                    for c in chunks:  # V6 Pas de limite — tous les chunks dans le Wiki
                        wiki.write_chunk(
                            content=c["content"],
                            source=source_name,
                            chunk_id=hashlib.md5(c["content"].encode()).hexdigest()[:8],
                            date="",
                            tags=[c.get("title", ""), c.get("level", "")],
                        )
                except Exception:
                    pass  # WikiWriter est optionnel, ne pas casser l'ingestion
                logger.info(f"✅ Ingestion OK: {os.path.basename(filepath)}")

                # NURU V12+ : Extraction structurée LLM (doc_structured / cv_structured)
                # Ne pas casser l'ingestion si l'extraction échoue
                try:
                    source_name = os.path.basename(filepath)
                    if not self.rag.is_doc_indexed(source_name, file_hash):
                        meta = await extract_document(text, filepath)
                        if meta:
                            self.rag.save_doc_meta(
                                source=source_name,
                                file_hash=file_hash,
                                doc_type=meta.doc_type,
                                json_data=meta.structured_json,
                            )
                            if meta.doc_type == "CV" and not self.rag.is_cv_indexed(source_name, file_hash):
                                self.rag.save_cv(source_name, file_hash, meta.structured_json)
                except Exception as e:
                    logger.warning(f"⚠️ Extraction structurée échouée pour {os.path.basename(filepath)}: {e}")

        except asyncio.TimeoutError:
            logger.error(f"⏱ Timeout sur {filepath} (>30s).")
        except Exception as e:
            logger.error(f"❌ Erreur critique {filepath}: {e}")

    async def auto_index_loop(self):
        """Boucle d'auto-indexation V8+ : scan récursif et hachage."""
        dirs_to_index = [
            Path.home() / "Documents",
            Path.home() / "Desktop",
            Path.home() / "Downloads"
        ]
        
        while True:
            logger.info("Scan auto-indexation V8+ démarré...")
            for base_dir in dirs_to_index:
                if not base_dir.exists():
                    continue
                
                # Scan récursif
                for root, _, files in os.walk(base_dir):
                    for file in files:
                        if not any(file.lower().endswith(e) for e in SUPPORTED_EXTENSIONS):
                            continue
                        filepath = os.path.join(root, file)
                        await self.index_file(filepath)
                        await asyncio.sleep(0.1) # Petite pause pour laisser respirer le CPU
            
            logger.info("Cycle d'auto-indexation terminé. Prochain scan dans 1h.")
            await asyncio.sleep(3600) # 1 heure


def reset_index():
    """Vide l'index RAG existant (supprime les fichiers de base) et arrête l'auto-indexation."""
    import sqlite3
    db_paths = [
        str(Path.home() / ".nuru" / "nuru_brain.db"),
        str(Path.home() / ".nuru" / "rag_index.db"),
        str(Path.home() / ".nuru" / "memory_v9.db"),
    ]
    for db in db_paths:
        try:
            if Path(db).exists():
                os.remove(db)
                logger.info(f"🗑️ Base supprimée: {db}")
        except Exception as e:
            logger.warning(f"⚠️ Impossible de supprimer {db}: {e}")
    # Signale l'arrêt via EventBus (si NuruCore tourne)
    try:
        from src.core.events import EventBus
        bus = EventBus()
        bus.emit_sync("index_reset", {})
    except Exception:
        pass
    logger.info("Index RAG réinitialisé — redémarrage nécessaire")
